"""
Procesador de documentos compartido
"""
import io
import logging
from pathlib import Path
from typing import Optional

import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image

from .models import Document, DocumentType
from .utils import (
    detect_document_type,
    encode_file_to_base64,
    get_mime_type,
    get_file_size,
    validate_image_file,
    validate_pdf_file,
    format_file_size
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Procesador de documentos multi-formato"""
    
    def __init__(self):
        self.supported_types = [
            DocumentType.PDF,
            DocumentType.IMAGE,
            DocumentType.EXCEL,
            DocumentType.WORD,
            DocumentType.TEXT
        ]
    
    def process_document(self, file_path: str) -> Document:
        """
        Procesa un documento y lo convierte al formato adecuado
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            Objeto Document procesado
            
        Raises:
            ValueError: Si el archivo no existe o no es soportado
            Exception: Si hay error al procesar el archivo
        """
        # Validar que el archivo existe
        if not Path(file_path).exists():
            raise ValueError(f"El archivo no existe: {file_path}")
        
        # Detectar tipo de documento
        doc_type = detect_document_type(file_path)
        file_name = Path(file_path).name
        size_bytes = get_file_size(file_path)
        
        logger.info(f"Procesando {file_name} ({format_file_size(size_bytes)}) - Tipo: {doc_type}")
        
        # Validar tipo soportado
        if doc_type == DocumentType.UNKNOWN:
            raise ValueError(f"Tipo de archivo no soportado: {file_name}")
        
        # Crear documento base
        document = Document(
            file_path=file_path,
            file_name=file_name,
            document_type=doc_type,
            size_bytes=size_bytes,
            mime_type=get_mime_type(file_path)
        )
        
        # Procesar según el tipo
        try:
            if doc_type == DocumentType.PDF:
                self._process_pdf(document)
            elif doc_type == DocumentType.IMAGE:
                self._process_image(document)
            elif doc_type == DocumentType.EXCEL:
                self._process_excel(document)
            elif doc_type == DocumentType.WORD:
                self._process_word(document)
            elif doc_type == DocumentType.TEXT:
                self._process_text(document)
            
            logger.info(f"✓ Documento procesado exitosamente: {file_name}")
            return document
            
        except Exception as e:
            logger.error(f"Error procesando {file_name}: {str(e)}")
            raise
    
    def _process_pdf(self, document: Document) -> None:
        """
        Procesa un archivo PDF
        
        Para PDFs, Claude en Bedrock puede procesarlos directamente,
        así que los codificamos en base64
        """
        # Validar PDF
        is_valid, message = validate_pdf_file(document.file_path)
        if not is_valid:
            raise ValueError(message)
        
        # Codificar a base64 para envío directo a Bedrock
        document.base64_content = encode_file_to_base64(document.file_path)
        document.mime_type = "application/pdf"
        
        # Opcionalmente, extraer texto para referencia
        try:
            with pdfplumber.open(document.file_path) as pdf:
                text_content = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                document.content = "\n\n".join(text_content)
        except Exception as e:
            logger.warning(f"No se pudo extraer texto del PDF: {str(e)}")
    
    def _process_image(self, document: Document) -> None:
        """
        Procesa un archivo de imagen
        
        Las imágenes se envían directamente a Claude en base64
        """
        # Validar imagen
        is_valid, message = validate_image_file(document.file_path)
        if not is_valid:
            raise ValueError(message)
        
        # Validar que la imagen se puede abrir
        try:
            with Image.open(document.file_path) as img:
                logger.info(f"Imagen: {img.format} {img.size} {img.mode}")
        except Exception as e:
            raise ValueError(f"No se puede abrir la imagen: {str(e)}")
        
        # Codificar a base64
        document.base64_content = encode_file_to_base64(document.file_path)
        
        # Asegurar mime_type correcto
        if not document.mime_type:
            extension = Path(document.file_path).suffix.lower()
            mime_mapping = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.webp': 'image/webp'
            }
            document.mime_type = mime_mapping.get(extension, 'image/jpeg')
    
    def _process_excel(self, document: Document) -> None:
        """
        Procesa un archivo Excel
        
        Excel no es soportado nativamente por Claude, así que
        convertimos a texto/CSV
        """
        try:
            # Leer todas las hojas
            excel_file = pd.ExcelFile(document.file_path)
            all_sheets_text = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(document.file_path, sheet_name=sheet_name)
                
                # Convertir a texto legible
                sheet_text = f"=== Hoja: {sheet_name} ===\n\n"
                sheet_text += df.to_string(index=False)
                all_sheets_text.append(sheet_text)
            
            document.content = "\n\n".join(all_sheets_text)
            logger.info(f"Excel procesado: {len(excel_file.sheet_names)} hojas")
            
        except Exception as e:
            raise ValueError(f"Error procesando Excel: {str(e)}")
    
    def _process_word(self, document: Document) -> None:
        """
        Procesa un archivo Word
        
        Word no es soportado nativamente por Claude, así que
        extraemos el texto
        """
        try:
            doc = DocxDocument(document.file_path)
            
            # Extraer texto de todos los párrafos
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extraer texto de tablas
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                tables_text.append("\n".join(table_data))
            
            # Combinar todo
            all_text = paragraphs
            if tables_text:
                all_text.append("\n=== TABLAS ===\n")
                all_text.extend(tables_text)
            
            document.content = "\n\n".join(all_text)
            logger.info(f"Word procesado: {len(paragraphs)} párrafos, {len(doc.tables)} tablas")
            
        except Exception as e:
            raise ValueError(f"Error procesando Word: {str(e)}")
    
    def _process_text(self, document: Document) -> None:
        """
        Procesa un archivo de texto plano
        """
        try:
            with open(document.file_path, 'r', encoding='utf-8') as f:
                document.content = f.read()
        except UnicodeDecodeError:
            # Intentar con otra codificación
            try:
                with open(document.file_path, 'r', encoding='latin-1') as f:
                    document.content = f.read()
            except Exception as e:
                raise ValueError(f"Error leyendo archivo de texto: {str(e)}")
